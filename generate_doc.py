import docx
from docx.shared import Pt, Inches

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

doc.add_heading('Technical Documentation: Job Portal Management System', 0)

# Section 1
doc.add_heading('1. Project Overview', level=2)
doc.add_paragraph('The Job Portal Management System is a dynamic web-based application designed to bridge the gap between job seekers (candidates) and employers (recruiters). Built using Python\'s Flask framework and backed by a robust MySQL database, the platform provides a centralized hub where employers can post job vacancies and candidates can search, filter, and apply for roles. The system offers distinct user interfaces and dashboards tailored to the specific needs of candidates and recruiters.')

# Section 2
doc.add_heading('2. Problem Statement', level=2)
doc.add_paragraph('The traditional job-seeking process is often fragmented, relying on dispersed classifieds, unorganized social media posts, or direct company websites. This fragmentation makes it difficult for candidates to find relevant opportunities and for recruiters to manage applicant pipelines efficiently. There is a need for a unified, intuitive, and secure platform that streamlines the recruitment lifecycle, offering reliable job matching, application tracking, and profile management for both parties.')

# Section 3
doc.add_heading('3. Objectives', level=2)
doc.add_paragraph('• Centralize Recruitment: Provide a single platform connecting recruiters and candidates seamlessly.')
doc.add_paragraph('• Streamline Job Discovery: Enable candidates to find suitable jobs quickly using advanced search and filtering mechanisms.')
doc.add_paragraph('• Simplify Job Management: Allow recruiters to post, view, and delete job listings efficiently through a dedicated dashboard.')
doc.add_paragraph('• Enhance User Experience: Deliver a responsive, modern, and accessible user interface using Bootstrap 5.')
doc.add_paragraph('• Data Integrity and Security: Ensure safe handling of user data, applications, and authentication using a structured relational database (MySQL).')

# Section 4
doc.add_heading('4. Features', level=2)
doc.add_heading('For Candidates:', level=3)
doc.add_paragraph('• Registration & Authentication: Secure account creation and login.')
doc.add_paragraph('• Job Search & Filtering: Search jobs by keywords, location, and job type (Full Time, Remote, etc.).')
doc.add_paragraph('• Job Application: Apply for open positions with a single click.')
doc.add_paragraph('• Save Jobs: Bookmark jobs to review or apply for later.')
doc.add_paragraph('• Resume Upload: Upload resumes (PDF/DOCX) securely to the platform.')
doc.add_paragraph('• Candidate Dashboard: Track the status of submitted applications and view applied jobs history.')

doc.add_heading('For Recruiters:', level=3)
doc.add_paragraph('• Registration & Authentication: Secure employer account management.')
doc.add_paragraph('• Job Posting: Publish detailed job postings (Title, Salary, Location, Type, Skills Required).')
doc.add_paragraph('• Job Management: View active job postings, track the number of applicants per job, and delete outdated listings.')
doc.add_paragraph('• Employer Dashboard: Centralized view of total active jobs and total overall applicants.')

# Section 5
doc.add_heading('5. Technologies Used', level=2)
doc.add_paragraph('• Backend: Python 3.x, Flask (Web Framework)')
doc.add_paragraph('• Database: MySQL, pymysql (Database Driver)')
doc.add_paragraph('• Frontend: HTML5, CSS3, JavaScript, Bootstrap 5.3 (CSS Framework)')
doc.add_paragraph('• Templating Engine: Jinja2')
doc.add_paragraph('• Version Control: Git')

# Section 6
doc.add_heading('6. System Architecture', level=2)
doc.add_paragraph('The project follows a classic Model-View-Controller (MVC) architectural pattern adapted for Flask:')
doc.add_paragraph('• Model (Database Layer): MySQL handles data persistence. Raw SQL queries are executed via the pymysql cursor to interface with the database.')
doc.add_paragraph('• View (Presentation Layer): Jinja2 templates (.html) rendered with dynamic data passed from the backend.')
doc.add_paragraph('• Controller (Application Logic): app.py contains the routing mechanisms, business logic, authentication state management, and database query executions.')

# Section 7
doc.add_heading('7. Module Explanation', level=2)
doc.add_paragraph('1. Authentication Module: Manages user sessions. Handles secure registration and role-based login.')
doc.add_paragraph('2. Job Management Module: Handles the creation (POST), retrieval (GET), and deletion (DELETE) of job records. Includes dynamic filtering logic.')
doc.add_paragraph('3. Application Module: Records the relationship between a candidate and a job. Prevents duplicate applications.')
doc.add_paragraph('4. Dashboard Module: Renders personalized views based on the user role. Fetches real-time statistics.')

# Section 8
doc.add_heading('8. Database Schema Explanation', level=2)
doc.add_paragraph('The database (jobportal) consists of a normalized relational schema:')
doc.add_paragraph('• users: user_id (PK), fullname, email, password, role.')
doc.add_paragraph('• companies: company_id (PK), company_name, location, description, website.')
doc.add_paragraph('• jobs: job_id (PK), recruiter_id (FK), company_id (FK), title, skills_required, location, salary, job_type.')
doc.add_paragraph('• applications: application_id (PK), job_id (FK), user_id (FK), status, applied_at.')
doc.add_paragraph('• saved_jobs: saved_id (PK), job_id (FK), user_id (FK).')

# Section 9
doc.add_heading('9. Workflow', level=2)
doc.add_paragraph('1. A user arrives at the Home Page and can browse featured jobs or search by keyword.')
doc.add_paragraph('2. The user registers an account, selecting either Candidate or Recruiter role.')
doc.add_paragraph('3. If Recruiter: The user is redirected to the Employer Dashboard to post and manage jobs.')
doc.add_paragraph('4. If Candidate: The user searches/filters jobs and applies for them.')
doc.add_paragraph('5. The application is logged in the MySQL database and tracked via the Candidate Dashboard.')

# Section 10
doc.add_heading('10. API & Route Details', level=2)
doc.add_paragraph('• GET / : Renders homepage')
doc.add_paragraph('• GET /jobs : Job listing page with filtering')
doc.add_paragraph('• GET /job/<job_id> : Job details page')
doc.add_paragraph('• GET/POST /login & /register : Authentication routes')
doc.add_paragraph('• GET /candidate/dashboard & /recruiter/dashboard : Role-specific dashboards')
doc.add_paragraph('• POST /post_job : Inserts new job record')
doc.add_paragraph('• POST /apply/<job_id> : Records job application')
doc.add_paragraph('• POST /save_job/<job_id> : Bookmarks a job')
doc.add_paragraph('• POST /upload_resume : Uploads resume to file system')
doc.add_paragraph('• POST /delete_job/<job_id> : Deletes a job posting')

# Section 11
doc.add_heading('11. Installation Steps', level=2)
doc.add_paragraph('1. Clone the repository.')
doc.add_paragraph('2. Set up a Python Virtual Environment: python -m venv venv')
doc.add_paragraph('3. Install dependencies: pip install -r requirements.txt')
doc.add_paragraph('4. Set up MySQL database with jobportal schema and configure credentials in app.py')
doc.add_paragraph('5. Run the application: python app.py')
doc.add_paragraph('6. Access at http://127.0.0.1:5000')

# Section 12
doc.add_heading('12. Screenshots Placeholders', level=2)
doc.add_paragraph('[Insert Screenshot 1: Home Page & Search Bar]')
doc.add_paragraph('[Insert Screenshot 2: User Registration Form]')
doc.add_paragraph('[Insert Screenshot 3: Job Listing & Filter Page]')
doc.add_paragraph('[Insert Screenshot 4: Candidate Dashboard]')
doc.add_paragraph('[Insert Screenshot 5: Recruiter Dashboard & Post Job Modal]')

# Section 13
doc.add_heading('13. Challenges Faced', level=2)
doc.add_paragraph('• Schema Mapping & Data Integrity: Ensuring that job applications gracefully cascade or halt deletion when a recruiter attempts to delete a job.')
doc.add_paragraph('• Dynamic Form Filtering: Constructing dynamic SQL queries in Python that safely handle varying active filters while avoiding SQL injection vulnerabilities.')
doc.add_paragraph('• UI/UX Consistency: Integrating Bootstrap modals and managing state shifts seamlessly between user roles.')

# Section 14
doc.add_heading('14. Future Scope', level=2)
doc.add_paragraph('• Password Hashing (bcrypt/werkzeug.security)')
doc.add_paragraph('• Advanced Admin Panel')
doc.add_paragraph('• Email Notifications via SMTP')
doc.add_paragraph('• Resume Parsing via NLP')

# Section 15
doc.add_heading('15. Conclusion', level=2)
doc.add_paragraph('The Job Portal Management System successfully demonstrates the integration of a Python web framework with a relational database to solve a real-world problem. The application is lightweight, responsive, and provides a solid foundation for both employers and job seekers, serving as an excellent proof-of-concept for modern web development architectures.')

doc.add_page_break()
doc.add_heading('Project Diagrams (Mermaid.js Source)', 0)
doc.add_paragraph('NOTE: Microsoft Word cannot render Mermaid graphs directly. The following blocks contain the raw Mermaid.js syntax. You can paste these code blocks into the Mermaid Live Editor (https://mermaid.live) to generate the visual graphical images, and then simply paste the downloaded images back into this Word document in place of the text.')

doc.add_heading('1. Use Case Diagram', level=2)
usecase = '''flowchart LR
    Candidate([Candidate])
    Recruiter([Recruiter])
    
    subgraph Job Portal System
        UC1(Register / Login)
        UC2(Search & Filter Jobs)
        UC3(Apply for Jobs)
        UC4(Save Jobs for Later)
        UC5(Upload Resume)
        UC6(Post New Job)
        UC7(Manage / Delete Jobs)
        UC8(View Dashboard Analytics)
    end
    
    Candidate --> UC1
    Candidate --> UC2
    Candidate --> UC3
    Candidate --> UC4
    Candidate --> UC5
    Candidate --> UC8
    
    Recruiter --> UC1
    Recruiter --> UC6
    Recruiter --> UC7
    Recruiter --> UC8'''
doc.add_paragraph(usecase)

doc.add_heading('2. Entity-Relationship (ER) Diagram', level=2)
er = '''erDiagram
    USERS {
        int user_id PK
        varchar fullname
        varchar email
        varchar password
        varchar role
    }
    COMPANIES {
        int company_id PK
        varchar company_name
        varchar location
        text description
        varchar website
    }
    JOBS {
        int job_id PK
        int recruiter_id FK
        int company_id FK
        varchar title
        text skills_required
        varchar location
        varchar salary
        varchar job_type
    }
    APPLICATIONS {
        int application_id PK
        int job_id FK
        int user_id FK
        varchar status
        timestamp applied_at
    }
    SAVED_JOBS {
        int saved_id PK
        int job_id FK
        int user_id FK
    }

    USERS ||--o{ JOBS : "Posts (if Recruiter)"
    COMPANIES ||--o{ JOBS : "Has"
    USERS ||--o{ APPLICATIONS : "Submits (if Candidate)"
    JOBS ||--o{ APPLICATIONS : "Receives"
    USERS ||--o{ SAVED_JOBS : "Saves"
    JOBS ||--o{ SAVED_JOBS : "Is Saved In"'''
doc.add_paragraph(er)


doc.add_heading('3. Class Diagram', level=2)
classd = '''classDiagram
    class User {
        +int user_id
        +String fullname
        +String email
        +String password
        +String role
        +register()
        +login()
        +logout()
    }
    class Candidate {
        +searchJobs()
        +applyForJob(job_id)
        +saveJob(job_id)
        +uploadResume()
        +viewDashboard()
    }
    class Recruiter {
        +postJob()
        +deleteJob(job_id)
        +viewApplicants()
        +viewDashboard()
    }
    class Job {
        +int job_id
        +int recruiter_id
        +int company_id
        +String title
        +String skills_required
        +String location
        +String salary
        +String job_type
        +getDetails()
    }
    class Application {
        +int application_id
        +int job_id
        +int user_id
        +String status
        +DateTime applied_at
        +submit()
    }
    
    User <|-- Candidate
    User <|-- Recruiter
    Recruiter "1" --> "*" Job : Posts
    Candidate "1" --> "*" Application : Submits
    Job "1" --> "*" Application : Receives'''
doc.add_paragraph(classd)


doc.add_heading('4. Workflow Flowchart', level=2)
flow = '''flowchart TD
    Start([User Visits Website]) --> Auth{Has Account?}
    Auth -- No --> Reg[Register Account]
    Auth -- Yes --> Login[Login]
    Reg --> Login
    
    Login --> Role{Role?}
    
    Role -- Candidate --> CDashboard[Candidate Dashboard]
    CDashboard --> Search[Search & Filter Jobs]
    Search --> Action{Action?}
    Action -- Apply --> ApplyJob[Submit Application]
    Action -- Save --> SaveJob[Save Job to Profile]
    ApplyJob --> CDashboard
    SaveJob --> CDashboard
    
    Role -- Recruiter --> RDashboard[Recruiter Dashboard]
    RDashboard --> Post[Post New Job]
    RDashboard --> Manage[Manage / Delete Active Jobs]
    Post --> RDashboard
    Manage --> RDashboard'''
doc.add_paragraph(flow)


doc.add_heading('5. System Architecture Diagram', level=2)
arch = '''flowchart TD
    Client([Client / Web Browser])
    
    subgraph Frontend [Presentation Layer - Frontend]
        HTML[HTML5 & Jinja2 Templates]
        CSS[CSS3 & Bootstrap 5]
        JS[Vanilla JavaScript]
    end
    
    subgraph Backend [Application Layer - Flask Backend]
        App[app.py / Routing & Logic]
        Auth[Authentication & Session Module]
        JobMod[Job Processing Module]
    end
    
    subgraph Database [Data Layer - MySQL]
        DB[(MySQL Database)]
        Tables[Users, Jobs, Companies, Applications]
    end
    
    Client <-->|HTTP GET/POST Requests| HTML
    HTML <--> App
    App <--> Auth
    App <--> JobMod
    
    Auth <-->|SQL Queries via pymysql| DB
    JobMod <-->|SQL Queries via pymysql| DB
    
    DB --- Tables'''
doc.add_paragraph(arch)

doc.save('Job_Portal_Documentation.docx')
